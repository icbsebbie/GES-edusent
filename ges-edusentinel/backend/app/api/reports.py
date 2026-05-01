import io
import json
from fastapi import APIRouter, Depends, Response
from sqlalchemy.orm import Session
from docx import Document
from reportlab.pdfgen import canvas
from ..db import get_db
from ..models import DataSubmission, OrgUnit, User
from ..deps import get_current_user
from ..org_tree import is_in_subtree

router = APIRouter(prefix="/reports")

@router.get("/term.docx")
def term_report_docx(term_id: int, dataset: str, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    subs = db.query(DataSubmission).filter(DataSubmission.term_id == term_id, DataSubmission.dataset == dataset).all()
    subs = [s for s in subs if is_in_subtree(db, user.org_unit_id, s.org_unit_id)]

    doc = Document()
    doc.add_heading("GES EduSentinel Term Report", level=1)
    doc.add_paragraph(f"Term ID: {term_id} | Dataset: {dataset}")
    doc.add_paragraph(f"Scope Root OrgUnit: {user.org_unit_id}")

    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "OrgUnit"
    hdr[1].text = "Submitted At"
    hdr[2].text = "Payload (JSON)"

    for s in subs[:200]:  # safety cap
        ou = db.get(OrgUnit, s.org_unit_id)
        row = table.add_row().cells
        row[0].text = ou.name if ou else str(s.org_unit_id)
        row[1].text = s.submitted_at.isoformat()
        row[2].text = s.payload_json[:500]

    buf = io.BytesIO()
    doc.save(buf)
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="term_{term_id}_{dataset}.docx"'}
    )

@router.get("/term.pdf")
def term_report_pdf(term_id: int, dataset: str, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    subs = db.query(DataSubmission).filter(DataSubmission.term_id == term_id, DataSubmission.dataset == dataset).all()
    subs = [s for s in subs if is_in_subtree(db, user.org_unit_id, s.org_unit_id)]

    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, 800, "GES EduSentinel Term Report")
    c.setFont("Helvetica", 10)
    c.drawString(50, 780, f"Term ID: {term_id} | Dataset: {dataset} | Scope Root: {user.org_unit_id}")

    y = 750
    for s in subs[:40]:
        ou = db.get(OrgUnit, s.org_unit_id)
        line = f"{(ou.name if ou else s.org_unit_id)} | {s.submitted_at.isoformat()} | {s.payload_json[:80]}"
        c.drawString(50, y, line)
        y -= 16
        if y < 50:
            c.showPage()
            y = 800

    c.save()
    return Response(
        content=buf.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="term_{term_id}_{dataset}.pdf"'}
    )
